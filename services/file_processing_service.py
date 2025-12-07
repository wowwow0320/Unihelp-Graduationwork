import os
import re
import io
import json
import asyncio
from pathlib import Path
from typing import Optional, Tuple, List, Dict

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from dotenv import load_dotenv
from llama_cloud_services import LlamaParse
from pdf2docx import Converter
# import pdfplumber  # LlamaParse를 정답지로 사용하므로 더 이상 필요 없음

# LlamaParse에 전달할 파싱 지시어 (전체 내용)
PARSING_INSTRUCTION = (
    "본문과 표를 구분해 주세요.\n"
    "- 표의 첫 행이 여러 줄로 구성되어 있다면, 이를 헤더로 간주하고 병합해 하나의 헤더로 만들어주세요.\n"
    "- 표는 정확히 html 형식으로 변환해주세요.\n"
    "- 표의 헤더는 반드시 각 열에 맞춰 분리해주세요.\n"
    "- 줄 바꿈 태그(<br/>)는 절대 사용하지 말고, 여러 조건이 있는 셀은 슬래시(/) 또는 쉼표(,)로 구분해주세요.\n"
    "- 병합된 셀은 해당 열에 맞춰 반복 삽입해주세요.\n"
    "- 열 수가 불균형한 행은 무조건 헤더 열 수에 맞춰 정렬해주세요.\n"
    "- 요약이나 설명은 절대 포함하지 마세요.\n"
)

class FileProcessorService:
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        load_dotenv()
        api_key = os.environ.get("LLAMA_CLOUD_API_KEY")
        if not api_key:
            raise ValueError("LLAMA_CLOUD_API_KEY 환경 변수가 설정되지 않았습니다.")
        
        # LlamaParse 클라이언트 초기화 (전체 설정)
        self.llama_parser = LlamaParse(
            api_key=api_key,
            parse_mode="parse_page_with_agent",
            model="openai-gpt-4-1",
            high_res_ocr=True,
            adaptive_long_table=True,
            outlined_table_extraction=True,
            output_tables_as_HTML=True,
            markdown_table_multiline_header_separator="<br />",
            system_prompt_append=PARSING_INSTRUCTION,
            page_separator="\n\n—\n\n",
        )
        print("✅ LlamaParse 클라이언트가 성공적으로 초기화되었습니다.")

    # --- 1. 메인 파이프라인 오케스트레이터 ---

    async def process_full_pipeline(self, pdf_path: str) -> Tuple[str, str, str, str]:
        """
        [최종 하이브리드 파이프라인 (v3: LlamaParse 정답지)]
        1. [Async] LlamaParse: 텍스트(.md) 추출 + 페이지 맵(정답지) 생성
        2. [Executor] pdf2docx: .docx 파일 생성
        3. [Executor] docx_parser + Matcher: .docx와 LlamaParse 맵을 매칭해 .html, .txt 생성
        """
        print(f"🚀 전체 하이브리드 파이프라인 시작: {pdf_path}")
        pdf_path_obj = Path(pdf_path)
        loop = asyncio.get_event_loop()

        # 1. LlamaParse로 텍스트(.md)와 페이지 맵 추출 (Async)
        print("🔧 [1/3] LlamaParse로 텍스트(.md) 및 페이지 맵 생성 중...")
        llama_task = self._parse_text_and_create_page_map_with_llama(pdf_path_obj)
        
        # 2. DOCX 변환 (Sync 함수를 Async로 실행)
        print("🔧 [2/3] DOCX 변환 시작...")
        docx_path_task = loop.run_in_executor(
            None, self.convert_pdf_to_docx, pdf_path
        )
        
        # 병렬 작업 1, 2가 완료되기를 기다림
        (markdown_path, page_map), docx_path = await asyncio.gather(
            llama_task,
            docx_path_task
        )
        
        print(f"✅ [1/3] Markdown (텍스트) 및 페이지 맵 생성 완료. 총 {len(page_map)} 페이지.")
        print(f"✅ [2/3] DOCX 저장 완료: {docx_path}")

        # 3. DOCX 파싱 및 LlamaParse 맵과 페이지 번호 매칭 (Sync 함수를 Async로 실행)
        print("🔧 [3/3] DOCX 파싱 및 페이지 번호 매칭 시작...")
        if not page_map:
             print("⚠️ 페이지 맵이 비어있어 매칭을 건너뜁니다. 페이지 번호가 -1로 표시됩니다.")
             
        html_path, rag_text_path = await loop.run_in_executor(
            None, 
            self._extract_tables_with_docx_and_matching,
            docx_path,
            page_map,
            pdf_path_obj.name # 메타데이터용
        )
        
        print(f"✅ [3/3] HTML (테이블) 및 RAG-TXT (K-V) 저장 완료: {html_path}, {rag_text_path}")
        print(f"🎉 전체 파이프라인 완료.")
        
        return docx_path, markdown_path, html_path, rag_text_path

    # --- 2. 파이프라인 구성 요소 ---

    def convert_pdf_to_docx(self, pdf_path: str, start_page: int = 0, end_page: Optional[int] = None) -> str:
        """ [Task 2] PDF를 DOCX로 변환 (기존과 동일) """
        print(f"📄 DOCX 변환 시작: {pdf_path}")
        pdf_path_obj = Path(pdf_path)
        docx_path = pdf_path_obj.with_suffix(".docx")
        try:
            cv = Converter(str(pdf_path_obj))
            cv.convert(str(docx_path), start=start_page, end=end_page)
            cv.close()
            print(f"✅ DOCX 저장 완료: {docx_path}")
        except Exception as e:
            print(f"❌ DOCX 변환 실패: {e}")
        return str(docx_path)

    async def _parse_text_and_create_page_map_with_llama(self, pdf_path_obj: Path) -> Tuple[str, Dict[int, str]]:
        """
        [Task 1] LlamaParse를 사용해 텍스트(.md)와 페이지 맵(정답지)을 동시에 생성합니다.
        """
        output_md_path = pdf_path_obj.with_suffix(".md")
        page_map: Dict[int, str] = {}
        processed_pages: List[str] = []
        
        try:
            result = await self.llama_parser.aparse(str(pdf_path_obj))
            markdown_documents = result.get_markdown_documents(split_by_page=True)
            
            for doc in markdown_documents:
                page_num = doc.metadata.get("page_number", -1)
                page_content_with_tables = doc.text # 테이블이 포함된 원본 텍스트
                
                # 1. 페이지 맵(정답지) 생성
                if page_num != -1:
                    # LlamaParse가 파싱한 (HTML 테이블 포함) 텍스트를 정답지로 사용
                    page_map[page_num] = page_content_with_tables 
                
                # 2. .md 파일용 텍스트 생성 (테이블 제거)
                text_only = self._preprocess_text(page_content_with_tables)
                # processed_pages.append(f"\n{text_only}") # 원본 코드
                processed_pages.append(f"\n{text_only}") # 사용자님이 수정한 코드

            # .md 파일 쓰기
            final_markdown = "\n\n—\n\n".join(processed_pages)
            with open(output_md_path, "w", encoding="utf-8") as f_md:
                f_md.write(final_markdown)
                
        except Exception as e:
            print(f"❌ LlamaParse 실패: {e}")
            
        return str(output_md_path), page_map

    def _extract_tables_with_docx_and_matching(self, docx_path: str, page_map: Dict[int, str], pdf_name: str) -> Tuple[str, str]:
        """
        [Task 3] 'python-docx'로 테이블을 파싱하고 'LlamaParse 페이지 맵'과 매칭하여
        페이지 번호가 포함된 .html과 .txt를 생성합니다.
        """
        html_path = Path(docx_path).with_suffix(".html")
        rag_text_path = Path(docx_path).with_suffix(".txt")
        
        all_kv_lines = []
        all_html_tables = []
        
        try:
            doc = Document(docx_path)
            
            # 1. DOCX에서 모든 <p>와 <table>을 순서대로 추출 (기존 로직 동일)
            full_parts = []
            section_blocks_list = self._get_section_blocks(doc)
            for idx, (section, section_blocks) in enumerate(zip(doc.sections, section_blocks_list), 1):
                body_parts = self._process_body(section_blocks)
                full_parts.extend(body_parts)
            
            full_html = self._prettify_html("\n".join(full_parts))
            soup = BeautifulSoup(full_html, "html.parser")

            # 2. HTML을 순회하며 테이블과 제목 추출, K-V 생성, 페이지 매칭
            for table in soup.find_all("table"):
                table_html = str(table)
                
                # --- A. 제목 추출 ---
                title = "제목 없음"
                prev_p = table.find_previous("p")
                if prev_p and not self._is_footer_text(prev_p.get_text(strip=True)):
                    title = prev_p.get_text(strip=True)

                # --- B. K-V 데이터 생성 (pandas 활용) ---
                try:
                    df_list = pd.read_html(io.StringIO(table_html), header=0)
                    if not df_list: continue
                    
                    df = df_list[0]
                    
                    # (사용자님의 다중 헤더 로직)
                    if not df.empty and df.iloc[0].astype(str).str.contains('대 학').any():
                        new_header = df.iloc[0]
                        df = df[1:] # 실제 데이터만 남김
                        df.columns = [f"{str(col).split('.')[0]} {val}" if 'Unnamed' not in str(col) else val for col, val in new_header.items()]
                    
                    if df.empty: continue

                    # --- C. 페이지 번호 매칭 (더욱 견고하게) ---
                    anchor_text_1 = "" # 테이블의 첫 행, 두 번째 값 (고유할 가능성 높음)
                    anchor_text_2 = "" # 테이블의 첫 행, 마지막 값 (추가 검증용)
                    try:
                        anchor_text_1 = str(df.iloc[0, 1]).strip() # 예: "경영학과(주)"
                        anchor_text_2 = str(df.iloc[0, -1]).strip() # 예: "130"
                    except IndexError:
                        pass # 테이블이 비었거나 구조가 이상하면 앵커 없음

                    found_page = -1
                    if (anchor_text_1 or anchor_text_2) and page_map:
                         found_page = self._find_page_for_anchor(
                             page_map, title, anchor_text_1, anchor_text_2
                         )
                    
                    # all_html_tables.append(f"\n# {title}\n") # 원본 코드
                    all_html_tables.append(f"\n# {title}\n") # 사용자님이 수정한 코드
                    all_html_tables.append(table_html) # 원본 HTML(구조가 올바른) 저장

                    # --- D. K-V 문장 생성 ---
                    for _, row in df.iterrows():
                        row_data = ", ".join([
                            f"{col}: {val}" 
                            for col, val in row.items() 
                            if pd.notna(val) and str(val).strip()
                        ])
                        if not row_data: continue

                        sentence = f"제목: {title}, {row_data}"
                        
                        metadata = {
                            "source": pdf_name,
                            "page": found_page,
                            "type": "table_kv"
                        }
                        meta_json = json.dumps(metadata, ensure_ascii=False)
                        all_kv_lines.append(f"{meta_json} {sentence}")

                except Exception as e:
                    print(f"⚠️ 테이블 K-V 변환/매칭 오류 (건너뜁니다): {e}")

            # RAG-TXT 파일 (K-V + 메타데이터) 저장
            with open(rag_text_path, 'w', encoding='utf-8') as f_txt:
                f_txt.write('\n'.join(all_kv_lines))
                
            # HTML 파일 (테이블 시각화용) 저장
            with open(html_path, 'w', encoding='utf-8') as f_html:
                f_html.write('\n\n'.join(all_html_tables))

        except Exception as e:
            print(f"❌ DOCX 파싱 및 매칭 전체 프로세스 실패: {e}")
            
        return str(html_path), str(rag_text_path)

    def _normalize_text_for_matching(self, text: str) -> str:
        """매칭을 위해 공백, 줄바꿈 등을 정규화합니다."""
        if not text:
            return ""
        # 모든 공백 문자(스페이스, 탭, 줄바꿈)를 단일 스페이스로 변환
        return re.sub(r'\s+', ' ', text).strip()

    def _find_page_for_anchor(self, page_map: dict, title: str, anchor_1: str, anchor_2: str) -> int:
        """
        LlamaParse 페이지 맵에서 제목과 2개의 앵커 텍스트로 페이지 번호를 검색합니다.
        공백/줄바꿈 차이를 무시하기 위해 정규화를 수행합니다.
        """
        
        # 매칭을 위해 앵커 텍스트와 제목을 정규화합니다.
        norm_title = self._normalize_text_for_matching(title)
        norm_anchor_1 = self._normalize_text_for_matching(anchor_1)
        norm_anchor_2 = self._normalize_text_for_matching(anchor_2)

        # 1순위: 앵커 2개가 모두 일치 (가장 신뢰도 높음)
        if norm_anchor_1 and norm_anchor_2:
            for page_num, page_text in page_map.items():
                norm_page_text = self._normalize_text_for_matching(page_text)
                if norm_anchor_1 in norm_page_text and norm_anchor_2 in norm_page_text:
                    return page_num
        
        # 2순위: 앵커 1만 일치
        if norm_anchor_1:
            for page_num, page_text in page_map.items():
                norm_page_text = self._normalize_text_for_matching(page_text)
                if norm_anchor_1 in norm_page_text:
                    return page_num
                    
        # 3순위: 제목과 앵커 1이 일치
        if norm_title != "제목 없음" and norm_anchor_1:
            for page_num, page_text in page_map.items():
                norm_page_text = self._normalize_text_for_matching(page_text)
                if norm_title in norm_page_text and norm_anchor_1 in norm_page_text:
                    return page_num
        
        return -1 # 모두 실패

    # --- 3. 유틸리티 메서드 (LlamaParse용) ---

    def _preprocess_text(self, text: str) -> str:
        """ LlamaParse 결과물에서 HTML 테이블을 제거하고 텍스트를 정리합니다. """
        text_without_tables = re.sub(r"<table.*?</table>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = text_without_tables.replace('\r\n', '\n')
        text = re.sub(r'<br\s*/?>', ' / ', text)
        text = re.sub(r'(?<!\n)\n(?!\n)', ' ', text) # 단일 줄바꿈은 공백으로
        text = re.sub(r'\n{2,}', '\n\n', text) # 여러 줄바꿈은 단락으로
        text = re.sub(r' +', ' ', text)
        return text.strip()

    # --- 4. 유틸리티 메서드 (python-docx 파싱용 - 원본 복원) ---

    def _prettify_html(self, html: str) -> str:
        html = re.sub(r'>\s+<', '><', html)
        html = re.sub(r'(</[^>]+>)', r'\1\n', html)
        return html.strip()

    def _iter_block_items(self, element, doc):
        """ docx의 body 요소를 순회하는 제너레이터 """
        for child in element.iterchildren():
            if isinstance(child, CT_P): yield Paragraph(child, doc)
            elif isinstance(child, CT_Tbl): yield Table(child, doc)

    def _extract_textbox_texts(self, paragraph):
        """ 단락 내 텍스트 박스 텍스트 추출 """
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        texts = []
        for txbx in paragraph._element.findall('.//w:txbxContent', ns):
            for t in txbx.findall('.//w:t', ns):
                if t.text: texts.append(t.text)
        return texts

    def _get_section_blocks(self, doc):
        """ docx를 구역(section)별 블록 리스트로 분리 """
        body = doc.element.body
        blocks = list(self._iter_block_items(body, doc))
        section_boundaries, current_blocks = [], []
        for blk in blocks:
            current_blocks.append(blk)
            if isinstance(blk, Paragraph) and blk._element.find('.//w:sectPr', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}) is not None:
                section_boundaries.append(current_blocks)
                current_blocks = []
        if current_blocks: section_boundaries.append(current_blocks)
        return section_boundaries

    def _is_footer_text(self, text: str) -> bool:
        """ 머리글/바닥글의 페이지 번호 등을 필터링하기 위한 함수 """
        text = text.strip()
        if not text: return True
        if re.fullmatch(r"-?\s*\d{1,4}\s*-?", text): return True
        if re.fullmatch(r"(p\.?|page)\s*\d{1.4}", text, re.IGNORECASE): return True
        return False

    def _process_header_footer(self, section_part, label):
        """ (참고) 머리글/바닥글 처리 함수 (현재 메인 로직에선 사용 안 함) """
        parts = []
        for para in section_part.paragraphs:
            combined_text = para.text.strip()
            for tb_text in self._extract_textbox_texts(para):
                if tb_text not in combined_text: combined_text += tb_text
            if combined_text and not self._is_footer_text(combined_text):
                parts.append(f"<p class='header'>[{label}] {combined_text}</p>")
        for table in section_part.tables:
            html_table = "<table border='1' class='header'>\n"
            for row in table.rows:
                html_table += "<tr>" + "".join(f"<td>{'<br>'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())}</td>" for cell in row.cells) + "</tr>\n"
            html_table += "</table>\n"
            parts.append(html_table)
        return parts

    def _process_body(self, blocks):
        """ docx 본문 블록(단락, 테이블)을 HTML로 변환 """
        parts = []
        for block in blocks:
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if text and not self._is_footer_text(text): parts.append(f"<p>{text}</p>")
                tb_combined = "".join(self._extract_textbox_texts(block)).strip()
                if tb_combined and not self._is_footer_text(tb_combined): parts.append(f"<p>{tb_combined}</p>")
            elif isinstance(block, Table):
                html_table = "<table border='1'>\n"
                for row in block.rows:
                    html_table += "<tr>" + "".join(f"<td>{'<br>'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())}</td>" for cell in row.cells) + "</tr>\n"
                html_table += "</table>\n"
                parts.append(html_table)
        return parts